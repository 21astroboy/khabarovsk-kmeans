from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
REPORT = ROOT / "report" / "Landsat_Khabarovsk_Kmeans_Report.docx"
GITHUB_URL = "https://github.com/21astroboy/khabarovsk-kmeans"


def set_run(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, align=None, size=11, bold=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, size=14 if level == 1 else 12, bold=True, color="1F4D78")
    return p


def set_cell(cell, text, bold=False, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cells[i], value, size=9, align=align)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    doc.add_paragraph()
    return table


def add_picture(doc, path, caption, width=Inches(6.2)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run(r, size=10, italic=True)


def build() -> None:
    data = json.loads((OUTPUTS / "analysis_results.json").read_text(encoding="utf-8"))
    meta = data["metadata"]
    area_rows = data["area_table"]
    metrics = data["kmeans"]["metrics"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    for line in [
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        "НИТУ МИСИС",
        "Институт компьютерных наук",
    ]:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True)

    doc.add_paragraph()
    add_paragraph(doc, "Лабораторная работа", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    add_paragraph(
        doc,
        "по курсу «Обработка изображений и анализ спутниковых снимков по данным ДЗЗ»",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
    )
    add_paragraph(
        doc,
        "«Автоматическая классификация спутниковых снимков Landsat методом K-means»",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=14,
        bold=True,
    )
    doc.add_paragraph()
    add_paragraph(doc, "Вариант Б: Python", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_paragraph()
    add_paragraph(doc, "Выполнил: ФИО __________________________", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "Группа: _______________________________", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "Проверила: доцент кафедры АСУ Леменкова П. А.", align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()
    add_paragraph(doc, "Москва, 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    add_heading(doc, "Введение")
    add_paragraph(
        doc,
        "Цель работы — выявить основные типы земного покрова на территории Хабаровска и его окрестностей "
        "методом автоматической классификации мультиспектрального снимка Landsat. Для классификации "
        "использован алгоритм K-means, который группирует пиксели по сходству спектральных характеристик "
        "без заранее подготовленной обучающей выборки.",
    )
    add_paragraph(
        doc,
        "Основные задачи: подготовить каналы Landsat B2-B7, выполнить классификацию с числом кластеров "
        "k = 5-8, выбрать итоговое значение k, интерпретировать кластеры, оформить карту и рассчитать "
        "площади выделенных классов.",
    )

    add_heading(doc, "Описание исходных данных")
    data_rows = [
        ["Спутник", "Landsat 9 OLI/TIRS"],
        ["Продукт", "Landsat Collection 2 Level-2"],
        ["Сцена", data["scene_id"]],
        ["Дата съемки", meta["date_acquired"]],
        ["Территория", "Хабаровск и окрестности, фрагмент 120 x 120 км"],
        ["Path / Row", f"{meta['path']} / {meta['row']}"],
        ["Облачность", f"{meta['cloud_cover']}%"],
        ["Система координат", f"WGS 84 / UTM Zone {meta['utm_zone']} ({meta['crs']})"],
        ["Использованные каналы", "B2, B3, B4, B5, B6, B7"],
        ["Источник данных", "USGS EarthExplorer"],
    ]
    add_table(doc, ["Параметр", "Значение"], data_rows, widths=[Cm(5), Cm(11)])
    add_picture(doc, OUTPUTS / "natural_color.png", "Рисунок 1 — Исходный снимок Landsat-9 в натуральных цветах 4-3-2", Inches(6.4))

    add_heading(doc, "Методика обработки")
    add_paragraph(
        doc,
        "В Python были загружены каналы поверхностной отражательной способности B2-B7. Значения Level-2 "
        "переведены в отражательную способность с использованием масштабного коэффициента 0.0000275 и "
        "смещения -0.2. По слою QA_PIXEL исключены пиксели облаков, теней, снега, цирруса и NoData.",
    )
    add_paragraph(
        doc,
        "Для классификации массив каналов был преобразован в матрицу «пиксель × канал», затем признаки "
        "стандартизованы с помощью StandardScaler. Для k = 5, 6, 7, 8 рассчитаны WCSS и коэффициент "
        "силуэта. Наибольшее значение силуэта получено при k = 6, поэтому итоговая карта построена для "
        "шести классов.",
    )
    metric_rows = [[m["k"], f"{m['wcss']:.0f}", f"{m['silhouette']:.3f}"] for m in metrics]
    add_table(doc, ["k", "WCSS", "Коэффициент силуэта"], metric_rows, widths=[Cm(2), Cm(5), Cm(5)])

    add_heading(doc, "Результаты классификации")
    add_picture(doc, OUTPUTS / "classification_map.png", "Рисунок 2 — Карта автоматической классификации K-means", Inches(6.4))

    table_rows = []
    for row in area_rows:
        table_rows.append(
            [
                row["cluster"],
                row["name"],
                f"{row['area_km2']:.2f}",
                f"{row['share_percent']:.2f}",
                row["ndvi"],
            ]
        )
    add_table(
        doc,
        ["Кластер", "Предполагаемый тип земного покрова", "Площадь, км²", "Доля, %", "NDVI"],
        table_rows,
        widths=[Cm(2), Cm(7), Cm(3), Cm(2.5), Cm(2)],
    )

    add_heading(doc, "Анализ результатов")
    add_paragraph(
        doc,
        "На территории исследования преобладают растительные классы. Наиболее крупный класс — "
        "разреженная растительность и сельскохозяйственные угодья: 5692.26 км², или 43.31% площади. "
        "Плотная древесная растительность занимает 3877.66 км², или 29.50%. Такое распределение "
        "соответствует ландшафту окрестностей Хабаровска, где за пределами городской застройки широко "
        "представлены лесные массивы, поля и пойменные территории.",
    )
    add_paragraph(
        doc,
        "Водные объекты занимают 597.92 км², или 4.55%. Этот класс хорошо выделяется благодаря высоким "
        "значениям MNDWI и низкому отражению в ближнем инфракрасном диапазоне. На карте водный класс "
        "соответствует руслу Амура и протокам поймы.",
    )
    add_paragraph(
        doc,
        "Городская застройка занимает 367.61 км², или 2.80%. Промышленные зоны, дороги и открытый грунт "
        "занимают 1905.24 км², или 14.49%. Эти классы частично смешиваются с песчаными и открытыми "
        "пойменными участками, что типично для неконтролируемой спектральной классификации: алгоритм "
        "разделяет пиксели по спектральному сходству, а не по заранее заданным тематическим категориям.",
    )
    add_paragraph(
        doc,
        "В целом результат отражает основные типы земного покрова: водную гладь Амура, городскую и "
        "транспортную инфраструктуру Хабаровска, лесистые участки и сельскохозяйственные территории. "
        "Основные ограничения связаны с остаточными облачными пикселями и неоднозначностью спектрально "
        "похожих поверхностей, например открытого грунта, песчаных отмелей и застроенных территорий.",
    )

    add_heading(doc, "Выводы")
    conclusions = [
        "Снимок Landsat-9 по территории Хабаровска был подготовлен и классифицирован методом K-means.",
        "Оптимальным по коэффициенту силуэта среди k = 5-8 оказалось значение k = 6.",
        "Классификация позволила выделить воду, городскую застройку, промышленные/открытые территории, плотную и разреженную растительность.",
        "Построена тематическая карта с легендой, координатной сеткой, масштабом и стрелкой севера.",
        "Рассчитаны площади и доли всех выделенных классов.",
    ]
    for item in conclusions:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        run = p.add_run("• ")
        set_run(run)
        run = p.add_run(item)
        set_run(run)

    add_heading(doc, "Список источников")
    sources = [
        "USGS EarthExplorer. Landsat Collection 2 Level-2. URL: https://earthexplorer.usgs.gov/",
        "USGS. Landsat Collection 2 Level-2 Science Products. URL: https://www.usgs.gov/landsat-missions/landsat-collection-2-level-2-science-products",
        "Rasterio documentation. URL: https://rasterio.readthedocs.io/",
        "Scikit-learn documentation: KMeans clustering. URL: https://scikit-learn.org/stable/modules/clustering.html#k-means",
        "Matplotlib documentation. URL: https://matplotlib.org/",
    ]
    for src in sources:
        add_paragraph(doc, src)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Приложение А. Код обработки")
    add_paragraph(
        doc,
        "Полный воспроизводимый код опубликован в GitHub-репозитории: "
        f"{GITHUB_URL}. Основной файл обработки: src/landsat_kmeans_khabarovsk.py. "
        "Код выполняет загрузку каналов, маскирование облаков по QA_PIXEL, подготовку стека, подбор k, "
        "классификацию K-means, расчет площадей и экспорт карт. Ниже приведен листинг основного скрипта.",
    )
    code = (ROOT / "src" / "landsat_kmeans_khabarovsk.py").read_text(encoding="utf-8")
    for chunk_start in range(0, len(code), 3500):
        chunk = code[chunk_start : chunk_start + 3500]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(chunk)
        run.font.name = "Courier New"
        run.font.size = Pt(6)

    REPORT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    build()
